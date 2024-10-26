import pandas as pd
import sys
import argparse
from docx import Document
from docx.shared import Inches
from read_solution_files import *
from tables import *
from utils import *
from visualizations import *
from docx.shared import Pt
from datetime import datetime
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def save_doc(doc, week_name):
    doc_path = f'../../reports/{week_name}.docx'
    doc.save(doc_path)
    print(f"Summary report saved to '{doc_path}'.")

def create_summary_report(df, week_name, type='separate', addonDf=None):
    # type: separate or integrative
    doc = Document()

    # authors
    authors_paragraph = doc.add_paragraph()
    authors_run = authors_paragraph.add_run('Authors: Weronika Zawadzka 151943 Eliza Czaplicka 151963') 
    authors_run.font.size = Pt(8) 
    authors_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # date
    date_paragraph = doc.add_paragraph()
    date_run = date_paragraph.add_run(f'Date: {datetime.now().strftime("%B %d, %Y")}')
    date_run.font.size = Pt(8) 
    date_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # note
    paragraph = doc.add_paragraph()
    run = paragraph.add_run('Best solutions have been checked with the solution checker')
    run.font.size = Pt(8) 
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    doc.add_heading(f'Summary Report for {week_name}', level=1)
    p = doc.add_paragraph('Based on methods from ')
    add_hyperlink(p, 'our github project', f"https://github.com/wkzawadzka/evolutionary-computation/tree/master/src/java/src/main/java/evcomp/{week_name}")
    insertHR(p)

    # prepare sections
    doc.add_heading(f'Problem description', level=2)
    add_problem_description(doc)
    doc.add_heading(f'Pseudocode of implemented algorithms', level=2)
    doc.add_page_break()

    # create summary statistics tables
    dfTemp = pd.concat([df, addonDf], ignore_index=True) if addonDf is not None else df
    doc.add_heading(f'Summary performance of each method', level=2)
    p = doc.add_paragraph("")
    insertHR(p)
    if type == 'separate':
        for instance in dfTemp['instance'].unique():
            doc.add_heading(f'Instance: {instance}', level=3)
            variable_labels, summary_stats = create_summary_table(dfTemp, instance)

            for variable, summary_df in summary_stats.items():
                doc.add_heading(f'Summary for {variable_labels[variable]}', level=4)
                add_table(doc, summary_df)
            
            p = doc.add_paragraph("")
            insertHR(p)
    else:
        variable_labels, summary_stats = create_intergative_summary_table(dfTemp)
        for variable, summary_df in summary_stats.items():
            doc.add_heading(f'Summary for {variable_labels[variable]}', level=4)
            add_table(doc, summary_df, type='integrative')
        
        p = doc.add_paragraph("")
        insertHR(p)

    doc.add_page_break()

    # 2D visualisations
    doc.add_heading(f'2D visualisations of best solutions', level=2)
    p = doc.add_paragraph("")
    insertHR(p)
    for instance in df['instance'].unique():
        doc.add_heading(f'Instance: {instance}', level=3)
        generate_visualizations(doc, df, instance)
        p = doc.add_paragraph("")
        insertHR(p)
    doc.add_page_break()

    # printed solutions
    doc.add_heading(f'Best solutions, indices', level=2)
    p = doc.add_paragraph("")
    insertHR(p)
    for instance in df['instance'].unique():
        doc.add_heading(f'Instance: {instance}', level=3)
        print_solutions_with_lowest_fval(doc, df, instance)
    doc.add_page_break()

    # conclusions
    doc.add_heading(f'Conclusions', level=2)

    # save the document
    save_doc(doc, week_name)


def main(week_name, report_type, addon):
    print(f"Week Name: {week_name}")
    print(f"Report Type: {report_type}")
    print(f"Addon: {addon}")

    addonDf = None
    results = read_solution_files(week_name)
    if addon=="yes":
        addon_weeks = list_previous_weeks(week_name)
        addonDf = [] 
        for week in addon_weeks:
            addonDf.append(read_solution_files(week))
        addonDf = pd.concat(addonDf, ignore_index=True)
    
    create_summary_report(results, week_name, report_type, addonDf)
          

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description='Create a report based on the provided week name and options.')
    
    parser.add_argument('week_name', type=str, help='Name of the week (e.g., w3_local_search)')
    
    parser.add_argument('--type', type=str, default='separate', 
                        help='Type of report: separate/integrative (default: "separate")')
    
    parser.add_argument('--addon', type=str, default='no', 
                        help='Whether to include previous methods (default: "no")')

    args = parser.parse_args()
    main(args.week_name, args.type, args.addon)


