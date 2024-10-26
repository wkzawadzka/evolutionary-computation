import docx
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
import pandas as pd

def add_hyperlink(paragraph, text, url):
    # https://stackoverflow.com/questions/47666642/adding-an-hyperlink-in-msword-by-using-python-docx
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a new run object (a wrapper over a 'w:r' element)
    new_run = docx.text.run.Run(
        docx.oxml.shared.OxmlElement('w:r'), paragraph)
    new_run.text = text

    # Set the run's style to the builtin hyperlink style, defining it if necessary
    new_run.style = get_or_create_hyperlink_style(part.document)
    # Alternatively, set the run's formatting explicitly
    # new_run.font.color.rgb = docx.shared.RGBColor(0, 0, 255)
    # new_run.font.underline = True

    # Join all the xml elements together
    hyperlink.append(new_run._element)
    paragraph._p.append(hyperlink)
    return hyperlink

#This is only needed if you're using the builtin style above
def get_or_create_hyperlink_style(d):
    """If this document had no hyperlinks so far, the builtin
       Hyperlink style will likely be missing and we need to add it.
       There's no predefined value, different Word versions
       define it differently.
       This version is how Word 2019 defines it in the
       default theme, excluding a theme reference.
    """
    if "Hyperlink" not in d.styles:
        if "Default Character Font" not in d.styles:
            ds = d.styles.add_style("Default Character Font",
                                    docx.enum.style.WD_STYLE_TYPE.CHARACTER,
                                    True)
            ds.element.set(docx.oxml.shared.qn('w:default'), "1")
            ds.priority = 1
            ds.hidden = True
            ds.unhide_when_used = True
            del ds
        hs = d.styles.add_style("Hyperlink",
                                docx.enum.style.WD_STYLE_TYPE.CHARACTER,
                                True)
        hs.base_style = d.styles["Default Character Font"]
        hs.unhide_when_used = True
        hs.font.color.rgb = docx.shared.RGBColor(0x05, 0x63, 0xC1)
        hs.font.underline = True
        del hs

    return "Hyperlink"


def insertHR(paragraph):
    # https://stackoverflow.com/questions/39006878/python-docx-add-horizontal-line
    p = paragraph._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)

def add_problem_description(doc):
    desc = (
        "We are given three columns of integers with a row for each node. "
        "The first two columns contain x and y coordinates of the node positions in a plane. "
        "The third column contains node costs. The goal is to select exactly 50% of the nodes "
        "(if the number of nodes is odd we round the number of nodes to be selected up) and form a Hamiltonian cycle "
        "(closed path) through this set of nodes such that the sum of the total length of the path plus the total "
        "cost of the selected nodes is minimized.\n\n"
        "The distances between nodes are calculated as Euclidean distances rounded mathematically to integer values. "
        "The distance matrix should be calculated just after reading an instance, and then only the distance matrix "
        "(no nodes coordinates) should be accessed by optimization methods to allow instances defined only by distance matrices."
    )
    

    p = doc.add_paragraph(desc)

def add_table(doc, summary_df, type='separate'):
    num_methods = summary_df.shape[0]  # Number of methods (rows)
    num_instances = summary_df.shape[1]  # Number of instances (columns)

    # Create a table with the number of rows equal to the length of the DataFrame plus one for the header
    if type == 'separate':
        table = doc.add_table(rows=summary_df.shape[0] + 1, cols=len(summary_df.columns))
    else:
        table = doc.add_table(rows=num_methods + 1, cols=num_instances + 1)

    # Add the header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Method'
    for i, column in enumerate(summary_df.columns):
        if type=='separate':
            hdr_cells[i].text = str(column)
            run = hdr_cells[i].paragraphs[0].runs[0]
        else:
            hdr_cells[i + 1].text = str(column)
            run = hdr_cells[i + 1].paragraphs[0].runs[0]
        run.font.bold = True

    # Add the data rows
    if type=='separate':
        for i, row in summary_df.iterrows():
            for j, value in enumerate(row):
                table.cell(i + 1, j).text = str(value) if value is not None else ''
    else:
        num_instances = summary_df.shape[1]
        for i, method in enumerate(summary_df.index):
            table.cell(i + 1, 0).text = method
            for j in range(num_instances):
                value = summary_df.loc[method].values[j]
                table.cell(i + 1, j+1).text = str(value) if value is not None else ''  # +1 to shift over


def read_input(instance):
    df = pd.read_csv(f"../../data/input/TSP{instance}.csv", header=None, delimiter=";")
    df.columns = ['x', 'y', 'cost'] 
    return df