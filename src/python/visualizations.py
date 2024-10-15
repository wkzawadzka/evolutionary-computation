import pandas as pd
from utils import read_input
import matplotlib.pyplot as plt
from io import BytesIO
from docx.shared import Inches

def generate_visualizations(doc, df, instance):
    problem = read_input(instance)
    df_subset = df[df['instance'] == instance]
    grouped = df_subset.groupby('method')  # Group by method

    for method, group in grouped:
        doc.add_heading(f"Method: {method}", level=4)
        # best_sol -> lowest f_val
        min_fval_row = group.loc[group['f_val'].idxmin()]
        doc.add_paragraph(f"Obj.f. value: {min_fval_row['f_val']}")
        solution_indices = min_fval_row['solution']

        problem_subset = problem.iloc[solution_indices]

        # Coordinates for all nodes
        all_x_coords = problem['x'].values
        all_y_coords = problem['y'].values
        all_costs = problem['cost'].values  # Costs for coloring and sizing all nodes

        # Coordinates for subset nodes (solution path)
        subset_x_coords = problem_subset['x'].values
        subset_y_coords = problem_subset['y'].values
        subset_costs = problem_subset['cost'].values  # Get costs for coloring subset nodes

        # Color and size normalization for all nodes
        norm = plt.Normalize(all_costs.min(), all_costs.max())  # Normalize using all node costs
        all_node_colors = plt.cm.hot_r(norm(all_costs))  # Colors for all nodes
        all_sizes = 100 * (all_costs - all_costs.min()) / (all_costs.max() - all_costs.min()) + 10  # Size for all nodes

        # Color and size normalization for subset nodes (for clarity)
        subset_norm = plt.Normalize(subset_costs.min(), subset_costs.max())  # Normalize for subset nodes
        subset_sizes = 100 * (subset_costs - subset_costs.min()) / (subset_costs.max() - subset_costs.min()) + 10  # Size for subset nodes

        plt.figure(figsize=(10, 8))

        # Plot all nodes with their respective colors and sizes
        plt.scatter(all_x_coords, all_y_coords, c=all_node_colors, s=all_sizes, edgecolor='black', alpha=0.5)  # All nodes

        # Plot only the path between nodes in the subset
        for i in range(len(solution_indices) - 1):
            plt.plot([subset_x_coords[i], subset_x_coords[i + 1]], [subset_y_coords[i], subset_y_coords[i + 1]], 'gray', linewidth=1)
        plt.plot([subset_x_coords[-1], subset_x_coords[0]], [subset_y_coords[-1], subset_y_coords[0]], 'gray', linewidth=1)

        # Highlight the subset nodes (solution path)
        plt.scatter(subset_x_coords, subset_y_coords, c=plt.cm.hot_r(subset_norm(subset_costs)), s=subset_sizes, edgecolor='black', alpha=0.8)

        plt.title(f'Visualization for Method: {method}, Instance: {instance}')
        plt.xlabel('X Coordinate')
        plt.ylabel('Y Coordinate')
        plt.grid(True)

        # Add labels to subset nodes
        for i, (x, y) in enumerate(zip(subset_x_coords, subset_y_coords)):
            plt.annotate(f'{i}', (x, y), textcoords="offset points", xytext=(0, 10), ha='center')

        # Color bar (based on all nodes' costs)
        sm = plt.cm.ScalarMappable(cmap='hot_r', norm=norm)
        sm.set_array([])  # Required for colorbar to work
        cbar = plt.colorbar(sm, ax=plt.gca()) 
        cbar.set_label('Cost', rotation=270, labelpad=15)

        # Save to doc
        memfile = BytesIO()
        plt.savefig(memfile)
        doc.add_picture(memfile, width=Inches(5))
        doc.add_paragraph()
        memfile.close()
